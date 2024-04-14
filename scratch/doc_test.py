from docx import Document
import os
document= Document()
name="guru"
add="gg \nggg \nn"
til=document.add_heading("         SRI ABIRAMI SILKS",0)
document.add_paragraph(f'Name:   {name}',style='Heading 1')
document.add_paragraph(f"ADDRES : {add}")
table = document.add_table(rows=1, cols=3)
records=[[1, 'rru', 566.0], [2, '555', 999.0], [3, '999', 555.0]]
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Sno'
hdr_cells[1].text = 'SAREE NO'
hdr_cells[2].text = 'PRICES'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = str(desc)

document.add_page_break()

document.save('demo.docx')
os.open('demo.docx',flags=1)